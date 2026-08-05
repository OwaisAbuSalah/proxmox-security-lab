# -*- coding: utf-8 -*-
"""Build the bilingual Word report from content.py."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import content as C

NAVY = RGBColor(0x0B, 0x2A, 0x4A)
TEAL = RGBColor(0x0E, 0x7C, 0x86)
GREY = RGBColor(0x55, 0x5B, 0x66)

doc = Document()

# base styles
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def set_rtl(par):
    p = par._p
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi'); pPr.append(bidi)
    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def ar(par):
    """mark a paragraph/run as arabic RTL and set complex-script font size"""
    set_rtl(par)

def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hexfill)
    tcPr.append(sh)

def add_ar(text, size=11, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    ar(p)
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color
    # complex script sizing
    rpr = r._r.get_or_add_rPr()
    sz = OxmlElement('w:szCs'); sz.set(qn('w:val'), str(size*2)); rpr.append(sz)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_en(text, size=11, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def h(text_ar, text_en, level=1):
    hp = doc.add_heading(level=level)
    ar(hp)
    r = hp.add_run(text_ar)
    r.font.color.rgb = NAVY
    # english subtitle line
    ep = doc.add_paragraph()
    er = ep.add_run(text_en)
    er.italic = True; er.font.size = Pt(10.5); er.font.color.rgb = TEAL
    ep.paragraph_format.space_after = Pt(8)
    return hp

# ---------- COVER ----------
for _ in range(3): doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar(t); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = t.add_run(C.PROJECT["title_ar"]); rt.bold=True; rt.font.size=Pt(28); rt.font.color.rgb=NAVY
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt2 = t2.add_run(C.PROJECT["title_en"]); rt2.font.size=Pt(18); rt2.font.color.rgb=TEAL
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER; ar(st); st.alignment=WD_ALIGN_PARAGRAPH.CENTER
rst = st.add_run(C.PROJECT["subtitle_ar"]); rst.font.size=Pt(14); rst.font.color.rgb=GREY
st2 = doc.add_paragraph(); st2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rst2 = st2.add_run(C.PROJECT["subtitle_en"]); rst2.italic=True; rst2.font.size=Pt(11); rst2.font.color.rgb=GREY
for _ in range(6): doc.add_paragraph()
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'{C.PROJECT["author"]}  |  {C.PROJECT["date"]}\n').font.size=Pt(11)
mr = meta.add_run(C.PROJECT["repo"]); mr.font.size=Pt(10); mr.font.color.rgb=TEAL
doc.add_page_break()

# ---------- 1. INTRO ----------
h("1. المقدمة والنطاق", "1. Introduction & Scope", 1)
add_ar(C.INTRO["about_repo_ar"]); add_en(C.INTRO["about_repo_en"], color=GREY, space_after=10)
add_ar(C.INTRO["why_proxmox_ar"]); add_en(C.INTRO["why_proxmox_en"], color=GREY, space_after=10)
add_ar(C.INTRO["scope_ar"], bold=True); add_en(C.INTRO["scope_en"], italic=True, color=GREY, space_after=10)

# quick stats table
add_ar("أرقام المكتبة:", bold=True, space_after=4)
stats = [("817","مهارة / Skills"),("29","مجال / Domains"),("6","أطر / Frameworks"),("Apache 2.0","الرخصة / License")]
tb = doc.add_table(rows=1, cols=4); tb.style="Light Grid Accent 1"; tb.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(v,k) in enumerate(stats):
    c=tb.rows[0].cells[i]
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(v+"\n"); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=NAVY
    r2=p.add_run(k); r2.font.size=Pt(9); r2.font.color.rgb=GREY
doc.add_paragraph()

# ---------- 1b. AI AGENT METHODOLOGY ----------
doc.add_page_break()
h("2. منهجية التنفيذ عبر وكيل ذكاء اصطناعي", "2. AI-Agent Execution Methodology", 1)
add_ar(C.INTRO["agent_method_ar"]); add_en(C.INTRO["agent_method_en"], color=GREY, space_after=10)
add_ar(C.INTRO["tiers_ar"]); add_en(C.INTRO["tiers_en"], color=GREY, space_after=10)

add_ar("بنية كل مهارة في الريبو:", bold=True, space_after=4)
struct = [
    ("SKILL.md", "ترويسة YAML (الاسم، الوصف، الوسوم، ربط الأطر) + خطوات Markdown", "YAML frontmatter + Markdown workflow"),
    ("scripts/agent.py", "الوكيل القابل للتشغيل الذي ينفّذ المهارة فعلياً", "The runnable agent that executes the skill"),
    ("references/", "مراجع تقنية إضافية (API reference)", "Additional technical references"),
    (".claude-plugin/", "يسجّل الريبو كإضافة لـ Claude Code", "Registers the repo as a Claude Code plugin"),
]
stb = doc.add_table(rows=0, cols=2); stb.style="Light List Accent 1"
for f, dar, den in struct:
    c = stb.add_row().cells
    r0 = c[0].paragraphs[0].add_run(f); r0.bold=True; r0.font.size=Pt(9.5); r0.font.name="Consolas"
    p1 = c[1].paragraphs[0]; ar(p1)
    r1 = p1.add_run(dar); r1.font.size=Pt(9.5)
    p2 = c[1].add_paragraph(); r2 = p2.add_run(den); r2.italic=True; r2.font.size=Pt(8.5); r2.font.color.rgb=GREY
doc.add_paragraph()

# ---------- 2. WHY / ARCHITECTURE ----------
h("3. معمارية المختبر", "3. Lab Architecture", 1)
add_ar("يحاكي المختبر شبكة مؤسسة مصغّرة على Proxmox، حيث تعمل كل مهارة كطبقة دفاع (defense-in-depth):", space_after=6)
add_en("The lab simulates a miniature enterprise network on Proxmox, where each skill acts as a defense-in-depth layer:", italic=True, color=GREY, space_after=6)
arch_rows = [
    ("الطبقة / Layer","المكوّن / Component","المهارة / Skill"),
    ("الوصول / Access","Tailscale + TLS 1.3 reverse proxy","4, 5"),
    ("الحافة / Edge","pfSense firewall + VLANs","1, 2"),
    ("المراقبة / Monitoring","Suricata IDS  →  Splunk/Wazuh SIEM","3, 6, 7"),
    ("التصليب / Hardening","OpenVAS scans + SBOM analysis","8, 10"),
    ("الخداع / Deception","Ransomware canary files","9"),
]
at = doc.add_table(rows=0, cols=3); at.style="Light List Accent 1"
for i,row in enumerate(arch_rows):
    cells = at.add_row().cells
    for j,val in enumerate(row):
        pp=cells[j].paragraphs[0]; rr=pp.add_run(val)
        if i==0: rr.bold=True; rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(cells[j],"0B2A4A")
        rr.font.size=Pt(10)
doc.add_paragraph()

# ---------- 4. THE 10 SKILLS ----------
doc.add_page_break()
h("4. المهارات العشرة", "4. The Ten Skills", 1)
add_ar("لكل مهارة: التعريف، خطوات سير العمل، الربط بالأطر (MITRE ATT&CK و NIST CSF)، ودورها في المختبر.", space_after=10)

for s in C.SKILLS:
    h(f'4.{s["n"]}  {s["name_ar"]}', f'{s["name_en"]}  ·  [{s["domain"]}]', 2)
    add_ar("ما هي: " + s["what_ar"], space_after=4)
    add_en("What it is: " + s["what_en"], italic=True, color=GREY, space_after=8)

    # steps two columns via table
    add_ar("خطوات سير العمل / Workflow steps:", bold=True, space_after=4)
    stt = doc.add_table(rows=len(s["steps_ar"]), cols=2); stt.style="Table Grid"
    for i in range(len(s["steps_ar"])):
        ca, ce = stt.rows[i].cells
        pa=ca.paragraphs[0]; ar(pa); ra=pa.add_run(f'{i+1}. {s["steps_ar"][i]}'); ra.font.size=Pt(9.5)
        pe=ce.paragraphs[0]; re_=pe.add_run(f'{i+1}. {s["steps_en"][i]}'); re_.font.size=Pt(9.5); re_.font.color.rgb=GREY
    stt.columns[0].width=Inches(3.3); stt.columns[1].width=Inches(3.3)
    doc.add_paragraph()

    # mapping table
    mt = doc.add_table(rows=2, cols=2); mt.style="Light Grid Accent 1"
    mt.rows[0].cells[0].paragraphs[0].add_run("MITRE ATT&CK").bold=True
    mt.rows[0].cells[1].paragraphs[0].add_run("NIST CSF 2.0").bold=True
    mt.rows[1].cells[0].paragraphs[0].add_run(", ".join(s["attack"]) if s["attack"] else "n/a")
    mt.rows[1].cells[1].paragraphs[0].add_run(", ".join(s["nist"]))
    doc.add_paragraph()

    # lab role
    lp = add_ar("دورها في المختبر: " + s["lab_ar"], color=TEAL, space_after=4)
    add_en("Role in the lab: " + s["lab_en"], italic=True, color=TEAL, space_after=12)
    if s["n"] % 3 == 0:
        doc.add_page_break()

# ---------- 5. REAL AGENT EXECUTION RESULTS ----------
doc.add_page_break()
h("5. نتائج تنفيذ الوكلاء الفعلية", "5. Actual Agent Execution Results", 1)
add_ar("نُفِّذت المهارات العشر جميعها فعلياً بتاريخ 2026-08-02: المستوى الأول على جهاز المضيف، والمستوى الثاني داخل "
       "عقدة Proxmox VE 9.2.2 حقيقية. المخرجات الكاملة (21 ملف دليل) محفوظة في مجلد agent-lab/evidence/.",
       bold=True, space_after=4)
add_en("All ten skills were actually executed on 2026-08-02: Tier 1 on the host machine, Tier 2 inside a real Proxmox VE "
       "9.2.2 node. Full outputs (21 evidence files) are stored in agent-lab/evidence/.",
       italic=True, color=GREY, space_after=10)

for i, r in enumerate(C.AGENT_RUNS, 1):
    p = add_ar(f'5.{i}  {r["skill"]}', bold=True, color=NAVY, space_after=3)
    # command block
    cp = doc.add_paragraph()
    cr = cp.add_run(r["cmd"]); cr.font.name="Consolas"; cr.font.size=Pt(8.5); cr.font.color.rgb=RGBColor(0x22,0x22,0x22)
    cp.paragraph_format.left_indent=Inches(0.25); cp.paragraph_format.space_after=Pt(4)
    shp = cp._p.get_or_add_pPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),'F2F4F7'); shp.append(sh)
    add_ar("النتيجة: " + r["result_ar"], space_after=3)
    add_en("Result: " + r["result_en"], italic=True, color=GREY, space_after=3)
    ep = add_ar("الدليل: " + r["evidence"], size=9, color=TEAL, space_after=12)

# Verified lab facts table
doc.add_page_break()
h("5.14  حقائق المختبر المُتحقَّق منها", "5.14  Verified Lab Facts", 2)
add_ar("كل سطر في هذا الجدول تم التحقق منه بأمر فعلي على العقدة، ومخرجاته محفوظة في مجلد الأدلة:", space_after=3)
add_en("Every row below was verified by an actual command on the node, with output saved in the evidence folder:",
       italic=True, color=GREY, space_after=6)
t2 = doc.add_table(rows=1, cols=2); t2.style="Medium Shading 1 Accent 1"
for j,v in enumerate(["البند / Item","القيمة المُتحقَّقة / Verified value"]):
    t2.rows[0].cells[j].paragraphs[0].add_run(v).bold=True
for k, v in C.LAB_FACTS:
    c = t2.add_row().cells
    r0=c[0].paragraphs[0].add_run(k); r0.bold=True; r0.font.size=Pt(9)
    r1=c[1].paragraphs[0].add_run(v); r1.font.size=Pt(8.5); r1.font.name="Consolas"
doc.add_paragraph()

add_ar("ملاحظة على النزاهة العلمية:", bold=True, color=NAVY, space_after=3)
add_ar("وثّقنا في هذا التقرير المحاولات الفاشلة وأسبابها (خطأ صياغة sed في إعداد الـ VLAN، واستخدام بادئة IPSet "
       "بدل الأسماء المستعارة في الجدار الناري)، كما وثّقنا القاعدتين اللتين لم تُطلقا في Suricata، والبدائل المبرَّرة "
       "التي استخدمناها بدل pfSense و Splunk، والخطوة الوحيدة غير المكتملة (تفعيل Tailscale). "
       "هذه الشفافية جزء من منهجية العمل وليست نقصاً فيه.", space_after=3)
add_en("Integrity note: this report documents the failed attempts and their causes (a malformed sed directive during VLAN "
       "setup, and an IPSet prefix used instead of aliases in the firewall), the two Suricata rules that did not fire, the "
       "justified substitutions used in place of pfSense and Splunk, and the single incomplete step (Tailscale activation). "
       "This transparency is part of the methodology, not a shortfall in it.",
       italic=True, color=GREY, space_after=10)

# ---------- 6. LAB BUILD ----------
doc.add_page_break()
h("6. بناء مختبر Proxmox", "6. Building the Proxmox Lab", 1)
add_ar(C.LAB_BUILD["host_ar"], bold=True, space_after=3)
add_en(C.LAB_BUILD["host_en"], italic=True, color=GREY, space_after=8)
add_ar("خطوات التجهيز المنفّذة:", bold=True, space_after=4)
for i,(a,e) in enumerate(zip(C.LAB_BUILD["steps_ar"], C.LAB_BUILD["steps_en"]),1):
    p=doc.add_paragraph(); ar(p); p.add_run(f"{i}. {a}").font.size=Pt(10)
    p.paragraph_format.space_after=Pt(1)
    p2=doc.add_paragraph(); r=p2.add_run(f"{i}. {e}"); r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=GREY
    p2.paragraph_format.space_after=Pt(7)

add_ar("ملاحظة على الافتراضية المتداخلة (Nested Virtualization):", bold=True, color=NAVY, space_after=3)
add_ar("تشغيل Proxmox (وهو hypervisor) داخل VMware يتطلب تمرير إمكانيات المعالج الافتراضية للنظام الضيف عبر الخيار "
       "vhv.enable = TRUE. وبدون تعطيل Hyper-V على المضيف، يحجز ويندوز طبقة الافتراضية ويمنع VMware من تمريرها.", space_after=3)
add_en("Running Proxmox (itself a hypervisor) inside VMware requires passing CPU virtualization extensions to the guest via "
       "vhv.enable = TRUE. Without disabling Hyper-V on the host, Windows claims the virtualization layer and blocks VMware from passing it through.",
       italic=True, color=GREY, space_after=10)

# ---------- 7. FRAMEWORK MAPPING SUMMARY ----------
doc.add_page_break()
h("7. جدول ربط الأطر الشامل", "7. Consolidated Framework Mapping", 1)
ft = doc.add_table(rows=1, cols=4); ft.style="Medium Shading 1 Accent 1"
hdr = ["#","المهارة / Skill","MITRE ATT&CK","NIST CSF 2.0"]
for j,val in enumerate(hdr):
    ft.rows[0].cells[j].paragraphs[0].add_run(val).bold=True
for s in C.SKILLS:
    cells = ft.add_row().cells
    cells[0].paragraphs[0].add_run(str(s["n"]))
    cells[1].paragraphs[0].add_run(s["name_en"]).font.size=Pt(9)
    cells[2].paragraphs[0].add_run(", ".join(s["attack"]) if s["attack"] else "n/a").font.size=Pt(8.5)
    cells[3].paragraphs[0].add_run(", ".join(s["nist"])).font.size=Pt(8.5)
doc.add_paragraph()

# ---------- 5. ROADMAP ----------
h("8. خطة النشر المرحلية", "8. Phased Deployment Roadmap", 1)
add_ar("ترتيب التنفيذ العملي للمشروع:", space_after=6)
for ar_t, en_t, ar_d, en_d in C.ROADMAP:
    p = add_ar(ar_t, bold=True, color=NAVY, space_after=2)
    add_en(en_t, italic=True, color=TEAL, space_after=2)
    add_ar(ar_d, space_after=2); add_en(en_d, color=GREY, space_after=8)

# ---------- 6. ETHICS ----------
h("9. اعتبارات أمنية وأخلاقية", "9. Security & Ethical Considerations", 1)
eth_ar = [
    "جميع الأنشطة دفاعية (blue-team) وتعليمية داخل مختبر معزول تملكه.",
    "لا يُوجَّه أي فحص أو هجوم محاكى نحو أنظمة أو شبكات لا تملك الإذن باختبارها.",
    "أنظمة الهدف الضعيفة (Metasploitable وغيرها) تبقى معزولة خلف pfSense ولا تُكشف للإنترنت.",
    "لا تُفتح واجهة Proxmox (المنفذ 8006) على الإنترنت مطلقاً  الوصول عبر Tailscale فقط.",
    "أخذ نسخ احتياطية (backups) للـ VMs المهمة واستخدام كلمات مرور قوية للحساب root.",
]
eth_en = [
    "All activity is defensive (blue-team) and educational inside an isolated lab you own.",
    "No scan or simulated attack is directed at systems/networks you are not authorized to test.",
    "Vulnerable target systems (Metasploitable, etc.) stay isolated behind pfSense, never internet-exposed.",
    "The Proxmox UI (port 8006) is never exposed to the internet; access is via Tailscale only.",
    "Back up important VMs and use strong passwords for the root account.",
]
for a,e in zip(eth_ar, eth_en):
    p=doc.add_paragraph(style="List Bullet"); ar(p); p.add_run(a).font.size=Pt(10.5)
    p2=doc.add_paragraph(style="List Bullet"); r=p2.add_run(e); r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY

# ---------- 7. CONCLUSION + REFS ----------
h("10. الخلاصة والمراجع", "10. Conclusion & References", 1)
add_ar("تُظهر هذه المهارات العشرة كيف تتحول قاعدة معرفة منظّمة (agentskills.io) إلى مختبر أمني متكامل قائم على Proxmox: من عزل الشبكة، مروراً بالمراقبة والكشف، وصولاً إلى التصليب والخداع, وكلها مربوطة بأطر MITRE ATT&CK و NIST CSF لتوثيق التغطية الدفاعية.", space_after=10)
add_en("These ten skills show how a structured, agentskills.io knowledge base turns into a complete Proxmox-based security lab, from network isolation, through monitoring and detection, to hardening and deception, all mapped to MITRE ATT&CK and NIST CSF for documented defensive coverage.", italic=True, color=GREY, space_after=12)
add_ar("المراجع:", bold=True, space_after=4)
refs = [
    C.PROJECT["repo"],
    "MITRE ATT&CK, https://attack.mitre.org",
    "NIST Cybersecurity Framework 2.0, https://www.nist.gov/cyberframework",
    "agentskills.io standard",
    "Proxmox VE, https://www.proxmox.com",
]
for r in refs:
    p=doc.add_paragraph(style="List Bullet"); rr=p.add_run(r); rr.font.size=Pt(10); rr.font.color.rgb=GREY

out = r"C:\Users\HP\automation\proxmox-security-lab\Proxmox-Security-Lab-Report.docx"
doc.save(out)
print("saved", out)
